from docx import Document
from docx.shared import Cm, Inches, Pt, Mm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document()


def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def set_row_height(row, height):
    row.height = height
    for cell in row.cells:
        cell.height = height


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def change_font_style(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)


def center(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = 1


def make_grey(table, *rows):
    for row in rows:
        table.rows[row._index].cells[0]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="808080"/>'.format(nsdecls('w'))))


p = document.add_paragraph('Ўзбекистон Республикаси Ички ишлар вазирининг шахсий таркиб бўйича буйруғи\nЛ О Й И Ҳ А С И')
run = p.runs[0]
run.font.bold = True
run.font.name = 'Times New Roman'
paragraph_format = p.paragraph_format
paragraph_format.alignment = 1

table = document.add_table(rows=10, cols=5)
hdr_cells = table.rows[0]
hdr_cells.cells[0].text = '№'
hdr_cells.cells[1].text = 'Унвони, Ф.И.Ш. '
hdr_cells.cells[2].text = 'Фотосурат'
hdr_cells.cells[3].text = 'Эгаллаб турган лавозими'
hdr_cells.cells[4].text = 'Таклиф қилинаётган лавозим'
table.rows[1].cells[0].merge(table.rows[1].cells[4]).text = 'ТАЙИНЛАНМОҚДА'
table.rows[2].cells[0].merge(table.rows[2].cells[4]).text = 'ВАЗИРЛИК  МАРКАЗИЙ  АППАРАТИ '
table.cell(3, 0).text = '1'
p1 = table.rows[3].cells[1].paragraphs[0]
run = p1.add_run()
run.add_text('майор\nААААА\nААААА\nААААА')
run.font.bold = True
table.rows[3].cells[1].add_paragraph('1974 йилда Тошкент шаҳрида туғилган, ўзбек')
table.cell(3, 2).paragraphs[0].add_run().add_picture('/home/dadakhon/Pictures/picture.png', Inches(1.0))
table.cell(3, 3).text = '**** ҳозирги ишлаб турган лавозим'
table.cell(3, 4).text = 'Келгусида хизмат олиб бораётган лавозим\n\n\n\n\nМуқаддам ушбу лавозимда: ААА ААА ААА'
table.rows[4].cells[0].merge(table.rows[4].cells[4]).text = 'ҚОРАҚАЛПОҒИСТОН РЕСПУБЛИКАСИ ИИВ '
table.cell(5, 0).text = '2'
p1 = table.rows[5].cells[1].paragraphs[0]
run = p1.add_run()
run.add_text('подполковник\nААААА\nААААА\nААААА')
run.font.bold = True
table.rows[5].cells[1].add_paragraph('\n1982 йилда Фарғона вилояти, Марғилон шаҳрида туғилган, ўзбек')
table.cell(5, 2).paragraphs[0].add_run().add_picture('/home/dadakhon/Pictures/picture.png', Inches(1.0))
table.cell(5, 3).text = '**** ҳозирги ишлаб турган лавозим'
table.cell(5, 4).text = 'Келгусида хизмат олиб бораётган лавозим\n\n\n\nМуқаддам ушбу лавозимда: ААА ААА ААА'
table.cell(6, 0).text = '3'
p1 = table.rows[6].cells[1].paragraphs[0]
run = p1.add_run()
run.add_text('майор\nААААА\nААААА\nААААА')
run.font.bold = True
table.rows[6].cells[1].add_paragraph('\n1985 йилда Тошкент шаҳрида туғилган, ўзбек')
table.cell(6, 2).paragraphs[0].add_run().add_picture('/home/dadakhon/Pictures/picture.png', Inches(1.0))
table.cell(6, 3).text = '**** ҳозирги ишлаб турган лавозим'
table.cell(6, 4).text = 'Келгусида хизмат олиб бораётган лавозим\n\n\n\n\nМуқаддам ушбу лавозимда: ААА ААА ААА'
table.rows[7].cells[0].merge(table.rows[7].cells[4]).text = 'ҚОЛДИРИЛМОҚДА '
table.rows[8].cells[0].merge(table.rows[8].cells[4]).text = 'ВАЗИРЛИК  МАРКАЗИЙ  АППАРАТИ '
table.cell(9, 0).text = '4'
p1 = table.rows[9].cells[1].paragraphs[0]
run = p1.add_run()
run.add_text('майор\nААА\nААА\nААА')
run.font.bold = True
table.rows[9].cells[1].add_paragraph('\n1986 йилда Тошкент вилояти, Зангиота туманида туғилган, ўзбек')
table.cell(9, 2).paragraphs[0].add_run().add_picture('/home/dadakhon/Pictures/picture.png', Inches(1.0))
table.cell(9, 3).text = '**** ҳозирги ишлаб турган лавозим '
table.cell(9, 4).text = '***** ихтиёрида\n\n\n\n\n\n\n '
p1 = table.rows[9].cells[4].add_paragraph('Асос: ')
run = p1.runs[0]
run.font.bold = True

set_column_width(table.columns[0], Cm(0.8))
set_column_width(table.columns[1], Cm(4.1))
set_column_width(table.columns[2], Cm(3.0))
set_column_width(table.columns[3], Cm(9.9))
set_column_width(table.columns[4], Cm(9.9))

set_row_height(table.rows[0], Mm(0.1))
set_row_height(table.rows[1], Mm(0))
set_row_height(table.rows[2], Mm(0))
set_row_height(table.rows[4], Cm(0))
set_row_height(table.rows[7], Cm(0))
set_row_height(table.rows[8], Cm(0))


sections = document.sections
for section in sections:
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.orientation = 1
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

make_rows_bold(hdr_cells, table.rows[1], table.rows[2], table.rows[4], table.rows[7], table.rows[8])
for i in range(0, 10):
    change_font_style(table.rows[i])

center(hdr_cells, table.rows[1], table.rows[2], table.rows[4], table.rows[7], table.rows[8])
make_grey(table, table.rows[1], table.rows[2], table.rows[4], table.rows[7], table.rows[8])

table.rows[1].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
table.rows[2].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
table.rows[4].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
table.rows[7].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
table.rows[8].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


document.save('povestka.docx')
