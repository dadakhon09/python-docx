from docx import Document
from docx.shared import Inches, Cm, RGBColor, Pt

document = Document()


def resize_table(table):
    set_column_width(table.columns[1], Cm(0.5))
    set_column_width(table.columns[2], Cm(10))


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def change_font_style(*tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)


table1 = document.add_table(rows=7, cols=3)
hdr_cells1 = table1.rows[0]
hdr_cells1.cells[0].text = 'ТАЙИНЛАНСИН:'
hdr_cells1.cells[1].text = ' '
hdr_cells1.cells[2].text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
row2_cells = table1.rows[2].cells
row2_cells[0].text = ''
row2_cells[1].text = ''
row2_cells[2].text = 'полковник АААА ВВВВВ (Н-000000), штатларни қайта ташкил этилиши муносабати билан, 2018 йилнинг 31' \
                     'декабрь кунидан ЛЛЛЛЛЛ бошлиғи лавозимидан озод этилиб, ЛЛЛЛЛЛ бошлиғи лавозимига;'

row4_cells = table1.rows[4].cells
row4_cells[0].text = ''
row4_cells[1].text = ''
row4_cells[
    2].text = 'полковник АААА ВВВВВ ББББББ (А-000000), **** лавозимидан озод этилиб, махсус унвони ҳарбий унвонга' \
              ' тенглаштирилиб, ***** лавозимига;'
row6_cells = table1.rows[6].cells
row6_cells[0].text = ''
row6_cells[1].text = ''
row6_cells[2].text = 'полковник АААА ВВВВВ ББББББ (А-000000), **** лавозимидан озод этилиб, махсус унвони ҳарбий' \
                     ' унвонга тенглаштирилиб, ***** лавозимига;'
table1.autofit = False

make_rows_bold(hdr_cells1)

resize_table(table1)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table2 = document.add_table(1, 1)
table2.rows[0].cells[0].text = '«Ўзбекистон Республикаси фуқаролари томонидан ҳарбий хизматни ўташ тартиби тўғрисидаги ' \
                               'Низом»нинг 36-моддаси, «е» бандига мувофиқ, резервдаги бор офицерлик унвони билан, ' \
                               'ихтиёрий равишда шартнома асосидаги ҳарбий хизматга чақирилиб,'

table3 = document.add_table(rows=10, cols=3)
row2_cells_t3 = table3.rows[1].cells
row2_cells_t3[0].text = 'ЮБОРИЛСИН:'
row2_cells_t3[2].text = 'Бухоро вилояти Пешку тумани мудофаа ишлари бўйича бўлимининг ҳарбий ҳисобида турган захирадаги' \
                        ' лейтенант ААААА ВВВВВ БББББ (А-000000), Қоровул қўшинлари қўмондони ихтиёрига;'
table3.rows[3].cells[2].text = 'Асос:'

table3.rows[5].cells[0].text = 'ЮБОРИЛСИН:'
table3.rows[5].cells[2].text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
table3.rows[7].cells[2].text = 'подполковник ААААА АААА ААААА (А-014296), штатларни қайта ташкил этилиши муносабати' \
                               ' билан, 2018 йилнинг 31 декабрь кунидан **** лавозимидан озод этилиб, келгуси хизмат' \
                               ' фаолиятини давом эттириш учун Жазони ижро этиш бош бошқармаси ихтиёрига;'
table3.rows[9].cells[2].text = 'Асос:'

resize_table(table3)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table4 = document.add_table(rows=3, cols=3)

table4.rows[0].cells[0].text = 'ТАСДИҚЛАНСИН:'
table4.rows[0].cells[2].text = 'САМАРҚАНД ВИЛОЯТИ ИЧКИ ИШЛАР БОШҚАРМАСИ БЎЙИЧА:'
table4.rows[2].cells[2].text = 'подполковник ААА ААА ААА (А-), **** бошлиғи лавозимида;'

resize_table(table4)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table5 = document.add_table(rows=3, cols=3)

table5.rows[0].cells[0].text = 'ҚОЛДИРИЛСИН:'
table5.rows[0].cells[
    2].text = 'подполковник ААА ААА ААА (А-), **** бошлиғи лавозимидан озод этилиб, бошқарма ихтиёрида.'
table5.rows[2].cells[2].text = 'Асос:'

resize_table(table5)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table6 = document.add_table(rows=7, cols=3)
hdr = table6.rows[0].cells[0].merge(table6.rows[0].cells[2])
hdr.text = 'Ўзбекистон Республикаси ИИВнинг 2018 йил 30 декабрдаги 2/4707-сонли кўрсатмаси талабларига мувофиқ, қуйидаги' \
           ' ходимлар 2019 йилнинг 14 январь кунидан 14 апрель кунига қадар бошланғич тайёргарликдан ўтишлари учун '

table6.rows[2].cells[0].text = 'ЮБОРИЛСИН:'
table6.rows[2].cells[2].text = 'ЎЗБЕКИСТОН РЕСПУБЛИКАСИ ИИВ АППАРАТИ БЎЙИЧА:'
row4 = table6.rows[4].cells[0].merge(table6.rows[4].cells[2])
row4.text = 'Вазирлик Ички ишлар органлари ходимларини бошланғич тайёрлаш ва малакасини ошириш маркази (Самарқанд шаҳар)' \
            ' ихтиёрига'
table6.cell(6, 2).text = '**** сафдор ходими сафдор ААА ААА ААА (А-);'

resize_table(table6)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table7 = document.add_table(rows=13, cols=3)
table7.cell(0, 0).text = 'ЎЗГАРТИРИШ КИРИТИЛСИН:'
table7.cell(0, 2).text = '**** сафдор ААА ААА АААнинг (А-) ҳисоб-тавсифловчи ҳужжатларига «АААА АААА АААА »-деб.'
table7.cell(2, 2).text = 'Асос:'
table7.cell(4, 0).text = 'БЕРИЛСИН:'
table7.cell(4, 2).text = 'Мирзо Улуғбек  номидаги  Ўзбекистон Миллий  университетининг сиртқи бўлимида таҳсил олаётган:'
table7.cell(6, 2).text = '**** сафдор ААА ААА ААА (А-), 2019 йилнинг 3 январь кунидан 30 январь кунига қадар тўловли ' \
                         'ўқув таътили;'
table7.cell(8, 2).text = 'Асос:'
table7.cell(10, 2).text = '*** подполковник ААА ААА АААга (А-), 2019 йилнинг 08 январь кунидан 2019 йилнинг 13 май' \
                          ' кунига қадар хомиладорлик ва туғиш таътили. \n2019 йилнинг 14 май кунидан хизмат вазифасини ' \
                          'бажаришга киришиши белгилансин;'
table7.cell(12, 2).text = 'Асос:'

resize_table(table7)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table8 = document.add_table(rows=9, cols=3)
table8.cell(0, 0).text = 'ЮБОРИЛСИН:'
table8.cell(0, 2).text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
table8.cell(2, 2).text = 'подполковник ААА ААА ААА (А-00), 2019 йилнинг 9 январь кунидан вазирлик кадрларининг амалдаги' \
                         ' захирасидан чиқарилиб, вазирлик ихтиёрида бўлган деб ҳисобланиб, буйруқ имзоланган кундан' \
                         ' Қорақалпоғистон Республикаси Ички ишлар вазирлиги ихтиёрига;'
table8.cell(4, 2).text = 'Асос:'
table8.cell(6, 2).text = 'подполковник ААА ААА ААА (А-), **** лавозимидан озод этилиб, келгуси хизмат фаолиятини давом' \
                         ' эттириш учун Қашқадарё вилояти ички ишлар бошқармаси ихтиёрига;'
table8.cell(8, 2).text = 'Асос:'

resize_table(table8)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table9 = document.add_table(rows=10, cols=3)
table9.cell(0, 0).text = 'ҚЎЛЛАНИЛСИН:'
table9.cell(0, 2).text = 'сабаблар, **** майор ААА ААА АААга (А-) “Қаттиқ ҳайфсан” интизомий жазоси;'
table9.cell(2, 2).text = 'Асос:'
row4 = table9.rows[4].cells[0].merge(table9.rows[4].cells[2])
row4.text = '«Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги Низом талабларига мувофиқ'
table9.cell(6, 2).text = 'ҚУРОЛЛИ КУЧЛАР ЗАХИРАСИГА'
table9.cell(7, 0).text = 'БЎШАТИЛСИН:'
table9.cell(8, 0).text = '144-бандининг «к» кичик бандига мувофиқ (ички ишлар органи ходимининг шаънига путур етказувчи' \
                         ' хатти-ҳаракатлар содир этганлиги учун)'
table9.cell(8, 2).text = '**** бошлиғи (Самарқанд шаҳри) подполковник ААА ААА ААА (А-0).'
table9.cell(9, 2).text = 'Асос:'

resize_table(table9)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table10 = document.add_table(rows=10, cols=3)
table10.cell(0, 0).text = 'ТАЙИНЛАНСИН:'
table10.cell(0, 2).text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
table10.cell(2, 2).text = 'майор АААА ААА ААА  (А-), *** лавозимидан озод этилиб, *** лавозимига;'
table10.cell(4, 2).text = 'майор ААА ААА ААА (А-), *** бошлиғи лавозимидан озод этилиб, “майор” махсус унвони ҳарбий ' \
                          'унвонга тенглаштирилиб, *** лавозимига;'
table10.cell(6, 2).text = 'капитан АААА ААА А (А-), **** лавозимидан озод этилиб, ўша ернинг ўзида **** лавозимига;'
table10.cell(8, 2).text = 'Жазони ижро этиш бош бошқармасидан келган капитан ААА ААА ААА (А-), *** лавозимига;'

resize_table(table10)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table11 = document.add_table(rows=3, cols=3)
table11.cell(0, 0).text = 'ЮБОРИЛСИН:'
table11.cell(0, 2).text = 'полковник ААА ААА ААА (К-), штатлар қайта ташкил этилиши муносабати билан, 2018 йилнинг ' \
                          '10 сентрябрь кунидан *** лавозимидан озод этилиб, вазирлик ихтиёрида бўлган деб ҳисобланиб, ' \
                          'буйруқ имзоланган кундан келгуси хизмат фаолиятини давом эттириш учун вазирлик Академияси ихтиёрига;'
table11.cell(2, 2).text = 'Асос:'

resize_table(table11)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table12 = document.add_table(rows=9, cols=3)
table12.cell(0, 0).text = 'БЕРИЛСИН:'
table12.cell(0, 2).text = 'Россия Федерациясининг Қозон федерал университети Тошкент шаҳар филиалининг сиртқи бўлимида ' \
                          'таҳсил олаётган:'
table12.cell(2, 2).text = '**** сержант АА ААА ААА (А-0), 2019 йилнинг 21 январь кунидан 11 февраль кунига тўловли ўқув' \
                          ' таътили;'
table12.cell(4, 2).text = 'Асос:'
table12.cell(6, 2).text = '*** сафдор ААА ААА ААА (А-), 2018 йилнинг 12 январь кунидан 2020 йилнинг 20 ноябрь қадар ' \
                          'бола икки ёшга тўлгунча парваришлаш учун иш ҳақи сақланмаган ҳолда таътили.\n2020 йилнинг 21' \
                          ' ноябрь кунидан хизмат вазифасини бажаришга киришиши белгилансин.'
table12.cell(8, 2).text = 'Асос:'

resize_table(table12)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table13 = document.add_table(rows=3, cols=3)
table13.cell(0, 0).text = 'ҲИСОБЛАНСИН:'
table13.cell(0, 2).text = '*** сафдор ААА ААА ААА (А-), бола парваришлаш таътил муддати тугашидан олдин 2019 йилнинг ' \
                          '01 февраль кунидан хизмат вазифасини бажаришга киришган деб. '
table13.cell(2, 2).text = 'Асос:'

resize_table(table13)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table14 = document.add_table(rows=3, cols=3)
table14.cell(0, 0).text = 'ЭЪЛОН ҚИЛИНСИН:'
table14.cell(0, 2).text = 'сабаб **** майор ААА ААА ААА (А-) “Қаттиқ ҳайфсан”;'
table14.cell(2, 2).text = 'Асос:'

resize_table(table14)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table15 = document.add_table(rows=5, cols=3)
table15.cell(0, 0).text = 'ЮБОРИЛСИН:'
table15.cell(0, 2).text = '**** подполковник ААА ААА ААА (У-), «Кўз микрохирургияси» цикли бўйича малака ошириш учун,' \
                          ' 2019 йилнинг 04 февраль кунидан 02 март кунига қадар Тошкент врачлар малакасини ошириш институти ихтиёрига;'
table15.cell(2, 2).text = 'Асос:'

resize_table(table15)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table16 = document.add_table(rows=8, cols=3)
table16.rows[0].cells[0].merge(
    table16.rows[0].cells[2]).text = '«Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги ' \
                                     'Низом талабларига мувофиқ'
table16.rows[1].cells[0].merge(table16.rows[1].cells[2]).text = ''
table16.rows[2].cells[0].merge(table16.rows[2].cells[2]).text = 'ҚУРОЛЛИ  КУЧЛАР  РЕЗЕРВИГА'
table16.cell(4, 0).text = 'БЎШАТИЛСИН:'
table16.columns[0].cells[5].merge(table16.columns[0].cells[
                                      6]).text = '144-бандининг «а» кичик бандига мувофиқ (пенсия олиш ҳуқуқини берадиган хизмат муддатини ўтамаганлиги)'
table16.columns[2].cells[4].merge(table16.columns[2].cells[5]).text = '**** сафдор ААА ААА АААА (А-).'
table16.cell(6, 2).text = 'Асос:'
resize_table(table16)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table17 = document.add_table(rows=2, cols=2)
table17.cell(0, 0).text = 'ПАСАЙТИРИЛСИН:'
table17.cell(0, 1).text = 'ЖИЗЗАХ  ВИЛОЯТИ  ИИБ  БЎЙИЧА: \n**** капитан ААА ААА ААА (А-)махсус унвони бир поғона «катта лейтенант»гача.'
table17.cell(1, 1).text = 'Асос:'

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

p = document.add_paragraph(
    '4. Ўзбекистон Республикаси Президентининг 2017 йил 29 ноябрдаги ПҚ–3413-сон қарори билан тасдиқланган Ички ишлар органларида хизматни ўташ тартиби тўғрисидаги низомнинг 46-бандига мувофиқ, ')

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table18 = document.add_table(rows=2, cols=2)
table18.cell(0, 0).text = 'МАХСУС УНВОНИДАН \nМАҲРУМ ЭТИЛСИН:'
table18.cell(0, 1).text = 'капитан ААА ААА ААА (А-) '
table18.cell(1, 1).text = 'Асос:'

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table19 = document.add_table(rows=9, cols=3)
table19.rows[0].cells[0].merge(table19.rows[0].cells[
                                   2]).text = '   «Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги Низом талабларига мувофиқ'
table19.rows[1].cells[0].merge(table19.rows[1].cells[2]).text = ''
table19.rows[2].cells[0].merge(table19.rows[2].cells[2]).text = 'ҚУРОЛЛИ  КУЧЛАР  ЗАХИРАСИГА'
table19.rows[3].cells[0].merge(table19.rows[3].cells[2]).text = ''
table19.cell(4, 0).text = 'БЎШАТИЛСИН:'
table19.cell(4, 2).text = 'ТОШКЕНТ ШАҲАР ИЧКИ ИШЛАР БОШ БОШҚАРМАСИ БЎЙИЧА:'
table19.cell(6, 0).text = '144-бандининг «д» кичик бандига мувофиқ \n(касаллиги туфайли) '
table19.cell(6, 2).text = 'полковник ААА ААА ААА (И-), штатларни қайта ташкил этилиши муносабати билан 2019 йилнинг 7 январь кунидан *** озод этилиб, бош бошқарма ихтиёрида бўлган деб ҳисобланиб, буйруқ имзоланган кундан.'
table19.cell(8, 2).text = 'Асос:'
resize_table(table19)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table20 = document.add_table(rows=4, cols=3)
table20.rows[0].cells[0].merge(table20.rows[0].cells[2]).text = 'ИСТЕЪФОГА'
table20.cell(2, 0).text = '144-бандининг «д» кичик бандига мувофиқ \n(касаллиги туфайли)'
table20.cell(2, 2).text = '****** бошлиғи полковник ААА ААА ААА (П-).'
table20.cell(3, 2).text = 'Асос:'
resize_table(table20)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table21 = document.add_table(rows=7, cols=3)
table21.cell(0, 0).text = 'ҲИСОБЛАНСИН:'
table21.cell(0, 2).text = '*** катта сержант ААА ААА (А-), бола парваришлаш таътили муддати тугашидан олдин 2019 йилнинг 01 февраль кунидан хизмат вазифасини бажаришга киришган деб;'
table21.cell(2, 2).text = 'Асос:'
table21.cell(4, 2).text = '*** сафдор ААА ААА ААА (А-), бола парваришлаш таътили муддати тугашидан олдин 2019 йилнинг 28 январь кунидан хизмат вазифасини бажаришга киришган деб;'
table21.cell(6, 2).text = 'Асос:'

resize_table(table21)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table22 = document.add_table(rows=8, cols=3)
table22.rows[0].cells[0].merge(
    table22.rows[0].cells[2]).text = '«Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги ' \
                                     'Низом талабларига мувофиқ'
table22.rows[1].cells[0].merge(table22.rows[1].cells[2]).text = ''
table22.rows[2].cells[0].merge(table22.rows[2].cells[2]).text = 'ҚУРОЛЛИ  КУЧЛАР  РЕЗЕРВИГА'
table22.cell(4, 0).text = 'БЎШАТИЛСИН:'
table22.columns[0].cells[5].merge(table22.columns[0].cells[
                                      6]).text = '144-бандининг «в» кичик бандига мувофиқ (пенсия олиш ҳуқуқини берадиган хизмат муддатини ўтаганлиги)'
table22.columns[2].cells[4].merge(table22.columns[2].cells[5]).text = '*** катта сержант ААА ААА ААА (А-);'
table22.cell(6, 2).text = 'Асос:'
resize_table(table22)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()


table23 = document.add_table(rows=6, cols=3)
table23.rows[0].cells[0].merge(table23.rows[0].cells[2]).text = '“Ўзбекистон Республикаси фуқаролари томонидан ҳарбий хизматни ўташ тартиби тўғрисида”ги Низомнинг 123-моддасига мувофиқ, ҳарбий хизматдаги узлуксиз хизмати учун бир йўла тўланадиган пул мукофоти'
table23.cell(2, 0).text = 'ТЎЛАНСИН:'
table23.cell(2, 2).text = 'ҚОРОВУЛ  ҚЎШИНЛАРИ  БОШ  БОШҚАРМАСИ  БЎЙИЧА:'
table23.cell(4, 2).text = '*** генерал-майор ААА ААА АААга (С-) 2017 йилнинг 25 январь кунидан 2019 йилнинг 25 январь кунига қадар, узлуксиз ўтаган хизмат йиллари учун икки ойлик амал маоши миқдорида.'
resize_table(table23)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()


table24 = document.add_table(rows=8, cols=3)
table24.rows[0].cells[0].merge(
    table24.rows[0].cells[2]).text = '«Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги Низом талабларига мувофиқ'
table24.rows[1].cells[0].merge(table24.rows[1].cells[2]).text = 'ИСТЕЪФОГА'
table24.rows[2].cells[0].merge(table24.rows[2].cells[2]).text = ' '
table24.cell(3, 0).text = 'БЎШАТИЛСИН:'
table24.cell(3, 2).text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
table24.columns[0].cells[5].merge(table24.columns[0].cells[
                                      6]).text = '144-бандининг «д» кичик бандига мувофиқ (касаллиги туфайли)'
table24.cell(5, 2).text = 'Вазирлик ихтиёридаги полковник ААА ААА ААА (П-).'
table24.cell(6, 2).text = 'Асос:'

resize_table(table24)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()


table25 = document.add_table(rows=2, cols=2)
table25.cell(0, 0).text = 'ҚАЙТА ТИКЛАНСИН:'
table25.cell(0, 1).text = 'АНДИЖОН  ВИЛОЯТИ  ИИБ  БЎЙИЧА: \n*** катта лейтенант ААА ААА АААнинг Ўзбекистон Республикаси Ички ишлар вазирининг 000 йил 00 ноябрдаги шахсий таркиб бўйича 000-сонли буйруғи билан бир поғона пасайтирилган «капитан» махсус унвони'
table25.cell(1, 1).text = 'Асос:'

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table26 = document.add_table(rows=5, cols=3)
table26.cell(0, 0).text = 'ЮБОРИЛСИН:'
table26.cell(0, 2).text = 'БУХОРО ВИЛОЯТИ ИЧКИ ИШЛАР БОШҚАРМАСИ БЎЙИЧА:'
table26.cell(2, 2).text = 'майор ААА ААА ААА (А-0), 2019 йилнинг 1 март кунидан *** лавозимидан озод этилиб, келгуси хизмат фаолиятини давом эттириши учун Ўзбекистон Республикаси Миллий гвардияси қўмондони ихтиёрига;'
table26.cell(4, 2).text = 'Асос:'

resize_table(table26)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()


table27 = document.add_table(rows=7, cols=3)
table27.rows[0].cells[0].merge(table27.rows[0].cells[2]).text = '*****'
table27.cell(2, 0).text = 'ЎРНАТИЛСИН:'
table27.cell(2, 2).text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
table27.cell(4, 2).text = '*** тезкор вакили подполковник ААА ААА ААнинг (А-), ойлик лавозим маошига 10 (ўн) фоиз миқдорида шахсий устама ҳақ;'
table27.cell(6, 2).text = 'Асос:'
resize_table(table27)

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()


table28 = document.add_table(rows=5, cols=3)
table28.cell(0, 0).text = 'ЎЧИРИЛСИН:'
table28.cell(0, 2).text = 'ҚОРАҚАЛПОҒИСТОН РЕСПУБЛИКАСИ ИИВ БЎЙИЧА:'
table28.cell(2, 2).text = '**** подполковник ААА ААА ААА (А-), вафот этганлиги сабабли 2019 йил 02 апрель кунидан шахсий таркиб рўйхатидан.\nУнинг ўлими хизмат мажбуриятларини бажариш билан боғлиқ деб ҳисоблансин.'
table28.cell(4, 2).text = 'Асос:'
resize_table(table28)


blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()


table29 = document.add_table(rows=9, cols=3)
table29.rows[0].cells[0].merge(
    table29.rows[0].cells[2]).text = '«Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги Низом талабларига мувофиқ'
table29.rows[1].cells[0].merge(table29.rows[1].cells[2]).text = ' '
table29.rows[2].cells[0].merge(table29.rows[2].cells[2]).text = 'ҚУРОЛЛИ КУЧЛАР ЗАХИРАСИГА'
table29.cell(4, 0).text = 'БЎШАТИЛСИН:'
table29.cell(4, 2).text = 'ТОШКЕНТ ШАҲАР ИЧКИ ИШЛАР БОШ БОШҚАРМАСИ БЎЙИЧА:'
table29.cell(6, 0).text = '144-бандининг «н» кичик бандига мувофиқ (бошқа турдаги давлат хизматига ўтганлиги, кадрларнинг амалдаги захирасига киритмасдан)'
table29.cell(6, 2).text = '**** подполковник АА АА АА (А-), 2019 йилнинг 12 апрель кунидан.\nЎзбекистон Республикаси Президентининг 2017 йил 05 майдаги ПФ-5037-сон Фармони асосида унинг Тадбиркорлик субъектларининг ҳуқуқлари ва қонуний манфаатларини ҳимоя қилиш бўйича вакил девонидаги иш вақти навбатдаги махсус унвон муддатига қўшилиши ҳамда уни белгиланган тартибда олиш кафолати сақланиб қолинсин.'
table29.cell(8, 2).text = 'Асос:'
resize_table(table29)

tables = [table1, table2, table3, table4, table5, table6, table7, table8, table9, table10, table11, table12, table13, table14, table15, table16, table17, table18, table19, table20, table21, table22, table23, table24, table25, table26, table27, table28, table29]
for i in range(0, 29):
    change_font_style(tables[i])

make_rows_bold(table1.rows[0], table2.rows[0], table3.rows[3], table3.rows[0], table3.rows[4], table4.rows[0], table5.rows[2], table6.rows[2], table6.rows[4], table6.rows[2], table7.rows[2], table7.rows[4], table7.rows[8], table7.rows[12], table8.rows[0], table8.rows[4], table8.rows[8], table9.rows[2])

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()
document.save('namuna.docx')
