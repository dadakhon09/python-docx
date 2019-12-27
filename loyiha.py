from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def resize_table(table):
    set_column_width(table.columns[1], Cm(0.5))
    set_column_width(table.columns[2], Cm(13))


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def change_font_style(*tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)


p = document.add_paragraph()
run = p.add_run()
p.alignment = 1
gerb = run.add_picture('/home/dadakhon/Pictures/gerb.png', width=Inches(1.05))

p = document.add_paragraph('ЎЗБЕКИСТОН РЕСПУБЛИКАСИ ИЧКИ ИШЛАР ВАЗИРИНИНГ\nБУЙРУҒИ')
run = p.runs[0]
p.alignment = 1
run.font.bold = True
run.font.name = 'Times New Roman'

p = document.add_paragraph('Шахсий таркиб бўйича')
run = p.runs[0]
p.alignment = 1
run.font.bold = True
run.font.name = 'Times New Roman'

table1 = document.add_table(rows=17, cols=3)
table1.cell(0, 0).text = 'ТАЙИНЛАНСИН:'
table1.cell(0, 2).text = 'ИЧКИ ИШЛАР ВАЗИРЛИГИ МАРКАЗИЙ АППАРАТИ БЎЙИЧА:'
table1.cell(2, 2).text = 'подполковник АААА АААА ААА (А-), **** лавозимидан озод этилиб, **** бошлиғи лавозимига;'
table1.cell(4, 0).text = 'ҚОЛДИРИЛСИН:'
table1.cell(4, 2).text = 'ҚОРАҚАЛПОҒИСТОН РЕСПУБЛИКАСИ ИИВ БЎЙИЧА:'
table1.cell(6, 2).text = 'подполковник ААА АААА ААА (А-0), *** лавозимидан озод этилиб, шу вазирлик ихтиёрида;'
table1.cell(8, 2).text = 'Асос:'
table1.rows[10].cells[0].merge(table1.rows[10].cells[
                                   2]).text = ' «Ички ишлар органларида хизматни ўташ тартиби тўғрисида»ги Низом талабларига мувофиқ '
table1.cell(12, 2).text = 'ҚУРОЛЛИ КУЧЛАР ЗАХИРАСИГА'
table1.cell(13, 0).text = 'БЎШАТИЛСИН:'
p = table1.rows[14].cells[0].paragraphs[0]
run = p.add_run()
run.add_text('144-бандининг «к» кичик бандига мувофиқ (ички ишлар органи ходимининг шаънига путур етказувчи хатти-ҳаракатлар содир этганлиги учун)')
run.font.bold = True
run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
p1 = table1.rows[14].cells[0].add_paragraph()
run1 = p1.add_run()
run1.add_text('топширган')
run1.font.bold = True
r = table1.rows[14].cells[2].paragraphs[0].add_run()
r.add_text('**** бошлиғи (Самарқанд шаҳри) подполковник ААА АААА АААА (А-).')
r.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
table1.cell(15, 2).text = 'Асос:'

resize_table(table1)


table1.rows[12].cells[2].paragraphs[0].alignment = 1

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table2 = document.add_table(rows=10, cols=2)
table2.cell(1, 0).text = 'Вазир'
table2.cell(2, 0).text = 'генерал-лейтенант'
table2.cell(2, 1).text = 'П.Р. Бобожонов'
table2.cell(5, 0).text = 'Тошкент шаҳри,'
table2.cell(6, 0).text = '2019  йил «_____» январь'
table2.cell(7, 0).text = '_____-сон'

table2.rows[1].cells[0].paragraphs[0].alignment = 1
table2.rows[2].cells[0].paragraphs[0].alignment = 1
table2.rows[5].cells[0].paragraphs[0].alignment = 1
table2.rows[6].cells[0].paragraphs[0].alignment = 1
table2.rows[7].cells[0].paragraphs[0].alignment = 1
table2.rows[2].cells[1].paragraphs[0].alignment = 2


set_column_width(table2.columns[0], Cm(5.1))
set_column_width(table2.columns[1], Cm(13.5))

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

table3 = document.add_table(1, 2)
table3.cell(0, 0).text = '****\nполковник'
table3.cell(0, 1).text = 'А.А.ААААА'

set_column_width(table3.columns[0], Cm(14.1))
set_column_width(table3.columns[1], Cm(4.5))

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

p = document.add_paragraph('К Е Л И Ш И Л Г А Н :')
run = p.runs[0]
p.alignment = 1
run.font.bold = True

table4 = document.add_table(rows=4, cols=2)
table4.cell(0, 0).text = '*****\nподполковник'
table4.cell(0, 1).text = 'А.А.ААААА'
table4.cell(1, 0).text = '*****\nподполковник'
table4.cell(1, 1).text = 'А.А.ААААА'
table4.cell(2, 0).text = '*****\nподполковник'
table4.cell(2, 1).text = 'А.А.ААААА'
table4.cell(3, 0).text = '*****\nподполковник'
table4.cell(3, 1).text = 'А.А.ААААА'

set_column_width(table4.columns[0], Cm(14.1))
set_column_width(table4.columns[1], Cm(4.5))

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

p = document.add_paragraph('Т А Қ С И М О Т :')
run = p.runs[0]
p.alignment = 1
run.font.bold = True

table5 = document.add_table(rows=16, cols=2)
table5.cell(0, 0).text = 'ТД Котибияти'
table5.cell(1, 0).text = 'КББ'
table5.cell(2, 0).text = 'МваМТТББ'
table5.cell(3, 0).text = 'Қорақалпоғистон Республикаси ИИВ'
table5.cell(4, 0).text = 'Андижон вилояти ИИБ'
table5.cell(5, 0).text = 'Бухоро вилояти ИИБ'
table5.cell(6, 0).text = 'Қашқадарё вилояти ИИБ'
table5.cell(7, 0).text = 'Навоий вилояти ИИБ'
table5.cell(8, 0).text = 'Наманган вилояти ИИБ'
table5.cell(9, 0).text = 'Самарқанд вилояти ИИБ'
table5.cell(10, 0).text = 'Тошкент шаҳар ИИББ '
table5.cell(11, 0).text = 'Фарғона вилояти ИИБ '
table5.cell(15, 0).text = 'Ж  А  М  И '

table5.rows[0].cells[1].paragraphs[0].add_run('1 нусха'), table5.rows[1].cells[1].paragraphs[0].add_run('1 нусха'),
table5.rows[2].cells[1].paragraphs[0].add_run('1 нусха'), table5.rows[3].cells[1].paragraphs[0].add_run('1 нусха'),
table5.rows[4].cells[1].paragraphs[0].add_run('1 нусха'), table5.rows[5].cells[1].paragraphs[0].add_run('1 нусха'),
table5.rows[6].cells[1].paragraphs[0].add_run('1 нусха'), table5.rows[7].cells[1].paragraphs[0].add_run('1 нусха'),
table5.rows[8].cells[1].paragraphs[0].add_run('1 нусха'), table5.rows[9].cells[1].paragraphs[0].add_run('1 нусха'),
table5.rows[10].cells[1].paragraphs[0].add_run('1 нусха'), table5.rows[11].cells[1].paragraphs[0].add_run('1 нусха'),
table5.rows[15].cells[1].paragraphs[0].add_run(' нусха')

for i in range(16):
    table5.rows[i].cells[1].paragraphs[0].alignment = 2

set_column_width(table5.columns[0], Cm(9.3))
set_column_width(table5.columns[1], Cm(9.3))

blank = document.add_paragraph(' ')
run_blank = blank.add_run()
run_blank.add_break()

p = document.add_paragraph('*****\nмайор                                          '
                           '                                                       '
                           '                    А.А.ААААА')
p.alignment = 0


for i in range(0, 8):
    blank = document.add_paragraph(' ')
    run_blank = blank.add_run()
    run_blank.add_break()

p = document.add_paragraph('Бажарди: *****\nкапитан А.А.ААААА тел.00-00')

sections = document.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

make_rows_bold(table1.rows[0], table1.rows[4], table1.rows[10], table1.rows[12], table1.rows[13], table1.rows[15],
               table2.rows[1], table2.rows[2], table5.rows[15])
tables = [table1, table2, table3, table4, table5]
for i in range(5):
    change_font_style(tables[i])

document.save('loyiha.docx')
